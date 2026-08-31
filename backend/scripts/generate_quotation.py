#!/usr/bin/env python3
"""
==============================================================================
 ARGUS — Python DOCX Renderer & PDF Converter for cPanel
 Called by Node.js backend (quotationController.js)
==============================================================================
"""

import sys
import json
import os
import subprocess

def render_docx(template_path, temp_docx_path, render_vars):
    """
    Renders variables into a DOCX template using docxtpl or python-docx fallback.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found at: {template_path}")

    # Try docxtpl (Jinja2 DOCX template engine)
    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(template_path)
        doc.render(render_vars)
        doc.save(temp_docx_path)
        return "docxtpl"
    except ImportError:
        pass

    # Fallback to python-docx directly if docxtpl is not installed
    try:
        from docx import Document
        doc = Document(template_path)
        
        # Simple text replacement across paragraphs and tables
        def replace_text_in_runs(paragraphs):
            for p in paragraphs:
                for key, val in render_vars.items():
                    placeholder = f"{{{key}}}"
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(val if val is not None else ''))

        replace_text_in_runs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_runs(cell.paragraphs)

        doc.save(temp_docx_path)
        return "python-docx"
    except ImportError:
        raise RuntimeError("Neither 'docxtpl' nor 'python-docx' is installed. Please install via pip: pip install docxtpl python-docx")


def convert_to_pdf(temp_docx_path, pdf_path):
    """
    Converts DOCX to PDF using LibreOffice CLI soffice, docx2pdf, or Mammoth+WeasyPrint fallback.
    """
    output_dir = os.path.dirname(pdf_path)
    os.makedirs(output_dir, exist_ok=True)

    # 1. Try LibreOffice CLI (soffice / libreoffice)
    soffice_bin = None
    for candidate in ["soffice", "libreoffice", "/usr/bin/soffice", "/usr/bin/libreoffice"]:
        try:
            res = subprocess.run([candidate, "--version"], capture_output=True, text=True)
            if res.returncode == 0:
                soffice_bin = candidate
                break
        except Exception:
            continue

    if soffice_bin:
        cmd = [soffice_bin, "--headless", "--convert-to", "pdf", temp_docx_path, "--outdir", output_dir]
        res = subprocess.run(cmd, capture_output=True, text=True)
        if res.returncode == 0 and os.path.exists(pdf_path):
            return f"LibreOffice CLI ({soffice_bin})"
        
        # If output filename was generated without timestamp prefix by soffice, rename it
        base_name = os.path.splitext(os.path.basename(temp_docx_path))[0]
        generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
        if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
            os.rename(generated_pdf, pdf_path)
            return f"LibreOffice CLI ({soffice_bin})"

    # 2. Try docx2pdf (Works on Windows/Mac with MS Word installed)
    try:
        from docx2pdf import convert
        convert(temp_docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return "docx2pdf"
    except Exception:
        pass

    # 3. Fallback: Mammoth (DOCX -> HTML) + WeasyPrint (HTML -> PDF)
    try:
        import mammoth
        from weasyprint import HTML

        with open(temp_docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value

        HTML(string=html_content).write_pdf(pdf_path)
        if os.path.exists(pdf_path):
            return "Mammoth + WeasyPrint"
    except Exception as e:
        pass

    raise RuntimeError("Failed to convert DOCX to PDF using LibreOffice, docx2pdf, or Mammoth+WeasyPrint.")


def main():
    if len(sys.argv) < 2:
        # Read payload from stdin if no CLI argument provided
        raw_input = sys.stdin.read().strip()
    else:
        # Check if first arg is a JSON string or file path
        first_arg = sys.argv[1]
        if os.path.exists(first_arg) and first_arg.endswith('.json'):
            with open(first_arg, 'r', encoding='utf-8') as f:
                raw_input = f.read().strip()
        else:
            raw_input = first_arg

    if not raw_input:
        sys.stderr.write("Error: No JSON payload provided to generate_quotation.py\n")
        sys.exit(1)

    try:
        data = json.loads(raw_input)
    except Exception as e:
        sys.stderr.write(f"Error parsing JSON payload: {e}\n")
        sys.exit(1)

    template_path = data.get("templatePath")
    temp_docx_path = data.get("tempDocxPath")
    pdf_path = data.get("pdfPath")
    render_vars = data.get("renderVars", {})

    if not template_path or not temp_docx_path or not pdf_path:
        sys.stderr.write("Error: Missing required paths (templatePath, tempDocxPath, pdfPath)\n")
        sys.exit(1)

    try:
        render_engine = render_docx(template_path, temp_docx_path, render_vars)
        converter_engine = convert_to_pdf(temp_docx_path, pdf_path)
        
        # Clean up temporary DOCX
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception:
                pass

        result_payload = {
            "success": True,
            "pdfPath": pdf_path,
            "renderEngine": render_engine,
            "converterEngine": converter_engine
        }
        print(json.dumps(result_payload))
        sys.exit(0)

    except Exception as err:
        # Clean up temporary DOCX on error
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except Exception:
                pass
        sys.stderr.write(f"Error: {str(err)}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
